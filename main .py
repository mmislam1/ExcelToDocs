


import xlrd
import docx
from docx import Document
from docx.shared import Inches
from datetime import datetime
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.table import _Cell
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx import Document
from docx.oxml.shared import OxmlElement, qn


def shade_cells(cells, shade):
    for cell in cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcVAlign = OxmlElement("w:shd")
        tcVAlign.set(qn("w:fill"), shade)
        tcPr.append(tcVAlign)

def set_cell_border(cell: _Cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
 
    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
 
    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
 
            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
 
            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

                    
#Given info                    
                    



excel_in=input("Enter excel file location: ")

#excel_in=str('./')+excel_in

excel_index=input("Enter excel sheet number: ")

excel_index=int(excel_index)




document = Document()

loc=(excel_in)

wb=xlrd.open_workbook(loc)
sheet=wb.sheet_by_index(excel_index)


name_is=str(sheet.cell(1,1))
name_is=name_is.replace(" ","_")
name_is=name_is.replace("'","")
file_out='Sajedur_'+ str(datetime.date(datetime.now()))+'x.docx'





key=0
count=0
table=[]
c=[]
calc=[]
for i in range(sheet.nrows):
    for j in range(sheet.ncols):
        
        
        if(sheet.cell_value(i,j)=='বেলা'):
            key=1
        if(key==1): 
            if(j==0 and sheet.cell_value(i,0)!=''):
                
                if(count==0):
                    count=i
                else:
                    x=i-count-1
                    c.append(x)
                    count=i
                
            if(sheet.cell_value(i,j)!='' and j!=3):
                table.append(sheet.cell_value(i,j))
        
        if(sheet.cell_value(i,j)=='নির্দেশনাবলী'):
            key=3
        
        if(sheet.cell_value(i,j)=='টোটাল'):
            x=i-count-1
            c.append(x)
            #calc.append(sheet.cell_value(i,j-1))
            key=2
            
        if(key==2 and sheet.cell_value(i,j)!=''):
            calc.append(sheet.cell_value(i,j))
        
            
#write to MS Word

def set_column_width(col, wid):
    for cell in col.cells:
        cell.width = wid

top_table= document.add_table(rows=2,cols=3)


top_table.style='Table Grid'

set_column_width(top_table.columns[0], Inches(4.5))
set_column_width(top_table.columns[1], Inches(3.5))
set_column_width(top_table.columns[2], Inches(1))

for row in top_table.rows:
    row.height = Inches(0.32)




a=top_table.cell(0,2)
b=top_table.cell(1,2)
A=a.merge(b)

paragraph = A.paragraphs[0]
run = paragraph.add_run()
run.add_picture('./logo.png', width = Inches(1), height = Inches(.6))


id_cells= top_table.rows[0].cells
id_cells[0].text= 'নামঃ '+ str(sheet.cell_value(1,1))
tarikh= datetime.date(datetime.now())
id_cells[1].text= 'তারিখঃ '+ str(tarikh)





id_cells2= top_table.rows[1].cells


id_cells2[0].text= str(calc[5])+'  ক্যালোরি,    রেশিওঃ  '+str(int(calc[10]))+':'+str(int(calc[11]))+':'+str(int(calc[12]))

id_cells2[1].text= 'টাইপঃ '


document.add_heading('',level=3)



#Main table starts from here 


header=document.add_table(rows=1,cols=7)
header.style='Table Grid'

#for row in header.rows:
#    row.height = Inches(0.33)

set_column_width(header.columns[0], Inches(1))
set_column_width(header.columns[1], Inches(2))
set_column_width(header.columns[2], Inches(1))
set_column_width(header.columns[3], Inches(1))
set_column_width(header.columns[4], Inches(1))
set_column_width(header.columns[5], Inches(1))
set_column_width(header.columns[6], Inches(1))

header_row=header.rows[0].cells

header_row[0].text='বেলা'
header_row[1].text='খাবার'
header_row[2].text='পরিমান'
header_row[3].text='ক্যালরি'
header_row[4].text='প্রোটিন'
header_row[5].text='ফ্যাট'
header_row[6].text='কার্ব'


for i in range(7):
    
    set_cell_border(
        header_row[i],
        top={"sz": 25, "val": "single", "color": "#000000", "space": "0"},
        bottom={"sz": 25, "val": "single", "color": "#000000", "space": "0"},
        start={"sz": 0, "val": "single", "color": "#000000", "space": "0"},
        end={"sz": 0, "val": "single", "color": "#000000", "space": "0"},
    )


table_index=7

c_cont=0

for c_x in range(len(c)) :
    
    
    if c_x==0:
        continue
    c_index=int(c[c_x])    
    
    body=document.add_table(rows=c_index,cols=7)
    body.style='Table Grid'
    

    set_column_width(body.columns[0], Inches(1))
    set_column_width(body.columns[1], Inches(2))
    set_column_width(body.columns[2], Inches(1))
    set_column_width(body.columns[3], Inches(1))
    set_column_width(body.columns[4], Inches(1))
    set_column_width(body.columns[5], Inches(1))
    set_column_width(body.columns[6], Inches(1))



    ax=int(c_index)
    ax-=1

    a=body.cell(0,0)
    b=body.cell(ax,0)
    A=a.merge(b)

    A.text=str(table[table_index])
    table_index +=1

    for i in range(int(c_index)):

        body_row=body.rows[i].cells

        for j in range(6):
            if j==1:
                body_row[j+1].text=str(table[table_index])+' '+ str(table[table_index+1])
                table_index+=2
            else:
                
                body_row[j+1].text=str(table[table_index])
                table_index+=1

    note=document.add_table(rows=1,cols=1)
    note.style='Table Grid'


    set_column_width(note.columns[0], Inches(7))
    note_row=note.rows[0].cells

    note_row[0].text='নোটঃ '
    
    shade_cells([note.cell(0, 0)], "#F9D0D0")
    
    set_cell_border(
        note_row[0],
        top={"sz": 15, "val": "single", "color": "#000000", "space": "0"},
        bottom={"sz": 15, "val": "single", "color": "#000000", "space": "0"},
        start={"sz": 0, "val": "single", "color": "#000000", "space": "0"},
        end={"sz": 0, "val": "single", "color": "#000000", "space": "0"},
    )


    
# Final calculation:

body=document.add_table(rows=3,cols=7)
body.style='Table Grid'
    

set_column_width(body.columns[0], Inches(1))
set_column_width(body.columns[1], Inches(2))
set_column_width(body.columns[2], Inches(1))
set_column_width(body.columns[3], Inches(1))
set_column_width(body.columns[4], Inches(1))
set_column_width(body.columns[5], Inches(1))
set_column_width(body.columns[6], Inches(1))

body_row=body.rows[0].cells 

body_row[0].text=str(calc[0])
body_row[1].text=''
body_row[2].text=''
body_row[3].text=''
body_row[4].text=str(round(calc[1], 2))
body_row[5].text=str(round(calc[2], 2))
body_row[6].text=str(round(calc[3], 2))


body_row=body.rows[1].cells 

body_row[0].text=str(calc[4])
body_row[1].text=''
body_row[2].text=''
body_row[3].text=str(round(calc[5], 2))
body_row[4].text=str(round(calc[6], 2))
body_row[5].text=str(round(calc[7], 2))
body_row[6].text=str(round(calc[8], 2))


body_row=body.rows[2].cells 

body_row[0].text=str(calc[9])
body_row[1].text=''
body_row[2].text=''
body_row[3].text=''
body_row[4].text=str(round(calc[10], 2))
body_row[5].text=str(round(calc[11], 2))
body_row[6].text=str(round(calc[12], 2))

# desktop/Chart-Box.xlsx


document.add_page_break()

document.save(file_out)



